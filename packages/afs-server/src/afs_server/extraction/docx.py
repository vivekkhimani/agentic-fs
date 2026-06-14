"""The ``docx`` rung — lightweight Word extraction via python-docx (pure-Python).

Word docs are reflowable (no fixed pages), so the whole document is one page of
markdown-ish text: paragraphs joined, table rows as pipe-separated lines. Fast
and reliable, no ML — common Office docs become readable synchronously.
"""

from __future__ import annotations

import asyncio
from typing import TYPE_CHECKING

from docx import Document

from afs_core.contracts import NormalizationError
from afs_core.models import NormalizedDocument, PageText, QualityReport

if TYPE_CHECKING:
    from afs_core.models import SourceDocument

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocxNormalizer:
    name = "docx"

    def accepts(self, doc: SourceDocument) -> bool:
        return (doc.content_type or "") == _DOCX_MIME or doc.local_path.suffix.lower() == ".docx"

    async def normalize(self, doc: SourceDocument) -> NormalizedDocument:
        return await asyncio.to_thread(self._extract_sync, str(doc.local_path))

    def _extract_sync(self, path: str) -> NormalizedDocument:
        document = Document(path)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        text = "\n\n".join(parts)
        if not text.strip():
            raise NormalizationError("empty_document", "no extractable text")
        return NormalizedDocument(
            pages=[PageText(number=1, markdown=text, source_locator="docx")],
            quality=QualityReport(page_count=1, char_count=len(text), min_chars_per_page=len(text)),
        )
