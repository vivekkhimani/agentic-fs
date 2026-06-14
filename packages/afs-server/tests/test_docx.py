"""The docx rung (python-docx) — conformance + paragraph/table extraction."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from afs_core.models import SourceDocument
from afs_core.testing import NormalizerConformance
from afs_server.extraction.docx import DocxNormalizer


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Quarterly Report", 0)
    doc.add_paragraph("Revenue grew across all regions.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "Revenue"
    table.rows[1].cells[0].text = "EMEA"
    table.rows[1].cells[1].text = "1.2M"
    doc.save(str(path))


def _src(path: Path) -> SourceDocument:
    return SourceDocument(filename=path.name, content_type=None, size=0, local_path=path)


class TestDocxNormalizer(NormalizerConformance):
    @pytest.fixture
    def normalizer(self) -> DocxNormalizer:
        return DocxNormalizer()

    @pytest.fixture
    def sample(self, tmp_path: Path) -> SourceDocument:
        p = tmp_path / "report.docx"
        _make_docx(p)
        return _src(p)


async def test_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    p = tmp_path / "r.docx"
    _make_docx(p)
    result = await DocxNormalizer().normalize(_src(p))
    md = result.pages[0].markdown
    assert "Quarterly Report" in md
    assert "Revenue grew across all regions." in md
    assert "EMEA | 1.2M" in md  # table row rendered as a pipe-separated line


def test_rejects_non_docx() -> None:
    assert DocxNormalizer().accepts(_src(Path("a.txt"))) is False
    assert DocxNormalizer().accepts(_src(Path("a.pdf"))) is False
